import PyPDF2
import docx
import pdfplumber
from fastapi import UploadFile, HTTPException
import re
from typing import Dict, Any, List, Tuple
import json
from io import BytesIO
import os
from transformers import pipeline, AutoTokenizer, AutoModelForTokenClassification
from transformers import AutoModelForSequenceClassification, AutoTokenizer as SeqTokenizer
import torch
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

try:
    import google.generativeai as genai
    GENAI_AVAILABLE = True
except ImportError:
    genai = None
    GENAI_AVAILABLE = False


class AdvancedResumeParser:
    """
    Enhanced resume parser using Hugging Face Transformers for advanced NLP capabilities.
    Combines traditional text extraction with AI-powered entity recognition and analysis.
    """

    def __init__(self):
        # Initialize Gemini for fallback
        self.api_key = os.getenv("GOOGLE_API_KEY")
        if self.api_key and GENAI_AVAILABLE and genai:
            genai.configure(api_key=self.api_key)
            self.gemini_available = True
        else:
            self.gemini_available = False

        # Initialize Hugging Face models
        self._initialize_models()

        # Skills dictionary for enhanced recognition
        self.skills_database = self._load_skills_database()

    def _initialize_models(self):
        """Initialize Hugging Face models for NER and other NLP tasks."""
        try:
            # Named Entity Recognition model
            self.ner_tokenizer = AutoTokenizer.from_pretrained("dbmdz/bert-large-cased-finetuned-conll03-english")
            self.ner_model = AutoModelForTokenClassification.from_pretrained("dbmdz/bert-large-cased-finetuned-conll03-english")
            self.ner_pipeline = pipeline("ner", model=self.ner_model, tokenizer=self.ner_tokenizer, aggregation_strategy="simple")

            # Sentence similarity model for skill matching
            self.similarity_tokenizer = AutoTokenizer.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")
            self.similarity_model = AutoModelForSequenceClassification.from_pretrained("sentence-transformers/all-MiniLM-L6-v2")

            print("✅ Advanced NLP models initialized successfully")
        except Exception as e:
            print(f"⚠️  Warning: Could not initialize advanced NLP models: {e}")
            print("Falling back to basic text processing")
            self.ner_pipeline = None

    def _load_skills_database(self) -> Dict[str, List[str]]:
        """Load comprehensive skills database for enhanced recognition."""
        return {
            "programming_languages": [
                "python", "java", "javascript", "typescript", "c++", "c#", "php", "ruby", "go", "rust",
                "swift", "kotlin", "scala", "perl", "r", "matlab", "shell", "bash", "powershell"
            ],
            "web_technologies": [
                "html", "css", "react", "angular", "vue", "node.js", "express", "django", "flask",
                "spring", "asp.net", "jquery", "bootstrap", "sass", "less", "webpack", "babel"
            ],
            "databases": [
                "mysql", "postgresql", "mongodb", "redis", "elasticsearch", "oracle", "sql server",
                "sqlite", "cassandra", "dynamodb", "firebase", "couchdb"
            ],
            "cloud_platforms": [
                "aws", "azure", "google cloud", "heroku", "digitalocean", "linode", "vercel",
                "netlify", "firebase", "cloudflare"
            ],
            "devops_tools": [
                "docker", "kubernetes", "jenkins", "gitlab ci", "github actions", "terraform",
                "ansible", "puppet", "chef", "nginx", "apache", "linux", "ubuntu", "centos"
            ],
            "data_science": [
                "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras", "jupyter",
                "matplotlib", "seaborn", "tableau", "power bi", "apache spark", "hadoop"
            ],
            "soft_skills": [
                "leadership", "communication", "problem solving", "teamwork", "project management",
                "agile", "scrum", "kanban", "time management", "analytical thinking"
            ]
        }

    async def parse(self, file: UploadFile) -> Dict[str, Any]:
        """Parse resume file and extract information using advanced NLP."""
        if not file or not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")

        # Extract raw text
        content = await self._extract_text(file)

        # Use advanced parsing pipeline
        return await self._analyze_with_advanced_nlp(content)

    async def _extract_text(self, file: UploadFile) -> str:
        """Extract text from PDF or DOCX file (same as original implementation)."""
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")

        try:
            content = await file.read()
            text = ""
            file_bytes = BytesIO(content)
            filename = file.filename.lower()

            if filename.endswith('.pdf'):
                try:
                    pdf_reader = PyPDF2.PdfReader(file_bytes)
                    text = " ".join(page.extract_text() for page in pdf_reader.pages)
                except Exception as e:
                    try:
                        file_bytes.seek(0)
                        with pdfplumber.open(file_bytes) as pdf:
                            text = " ".join(page.extract_text() for page in pdf.pages)
                    except Exception as pdf_e:
                        raise HTTPException(
                            status_code=400,
                            detail=f"Failed to read PDF file: {str(e)}. pdfplumber error: {str(pdf_e)}"
                        )

            elif filename.endswith(('.doc', '.docx')):
                try:
                    file_bytes.seek(0)
                    doc = docx.Document(file_bytes)
                    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                    if not paragraphs:
                        for table in doc.tables:
                            for row in table.rows:
                                paragraphs.extend(cell.text.strip() for cell in row.cells if cell.text.strip())

                    text = " ".join(paragraphs)

                    if not text.strip():
                        raise HTTPException(
                            status_code=400,
                            detail="No readable text found in the Word document. Please check if the document contains text content."
                        )
                except Exception as e:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Failed to read Word document: {str(e)}. Please ensure the document is not corrupted and is a valid .doc or .docx file."
                    )

            else:
                raise HTTPException(
                    status_code=400,
                    detail="Unsupported file format. Please upload a PDF or DOCX file."
                )

            if not text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="No text could be extracted from the file."
                )

            return text.strip()

        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process file: {str(e)}"
            )

    async def _analyze_with_advanced_nlp(self, text: str) -> Dict[str, Any]:
        """Analyze resume text using advanced NLP techniques."""
        try:
            # Extract entities using NER
            entities = self._extract_entities(text) if self.ner_pipeline else {}

            # Extract skills using pattern matching and ML
            skills = self._extract_skills(text)

            # Extract sections using text analysis
            sections = self._identify_sections(text)

            # Structure the information
            structured_data = {
                "personalInfo": self._extract_personal_info(text, entities),
                "workExperience": self._extract_work_experience(sections.get("experience", "")),
                "education": self._extract_education(sections.get("education", "")),
                "skills": skills,
                "highlights": self._extract_highlights(text)
            }

            # Fallback to Gemini if advanced parsing fails
            if not any(structured_data.values()):
                print("⚠️  Advanced NLP parsing yielded no results, falling back to Gemini")
                return await self._analyze_with_gemini(text)

            return structured_data

        except Exception as e:
            print(f"⚠️  Advanced NLP analysis failed: {e}")
            if self.gemini_available:
                print("Falling back to Gemini analysis")
                return await self._analyze_with_gemini(text)
            else:
                raise HTTPException(
                    status_code=500,
                    detail=f"Failed to analyze resume: {str(e)}"
                )

    def _extract_entities(self, text: str) -> Dict[str, List[str]]:
        """Extract named entities using Hugging Face NER model."""
        if not self.ner_pipeline:
            return {}

        try:
            entities = self.ner_pipeline(text)
            organized_entities = {
                "PERSON": [],
                "ORG": [],
                "GPE": [],  # Geopolitical entities (cities, countries)
                "DATE": []
            }

            for entity in entities:
                entity_type = entity["entity_group"]
                if entity_type in organized_entities:
                    if entity["word"] not in organized_entities[entity_type]:
                        organized_entities[entity_type].append(entity["word"])

            return organized_entities
        except Exception as e:
            print(f"⚠️  NER extraction failed: {e}")
            return {}

    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills using pattern matching and similarity matching."""
        found_skills = set()
        text_lower = text.lower()

        # Direct pattern matching
        for category, skill_list in self.skills_database.items():
            for skill in skill_list:
                if skill.lower() in text_lower:
                    found_skills.add(skill.title())

        # Extract additional skills using regex patterns
        skill_patterns = [
            r'\b(?:proficient|experienced|skilled|knowledgeable)\s+in\s+([A-Za-z\s+#]+)',
            r'\b([A-Za-z\s+#]+)\s+(?:development|programming|framework|library|tool|platform)',
            r'\b(?:using|with|via|through)\s+([A-Za-z\s+#]+)\s+(?:framework|library|tool)'
        ]

        for pattern in skill_patterns:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                skill = match.strip().title()
                if len(skill) > 2 and skill not in ['The', 'And', 'For', 'Are', 'But']:
                    found_skills.add(skill)

        return list(found_skills)

    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify different sections in the resume."""
        sections = {}
        text_lower = text.lower()

        # Common section headers
        section_patterns = {
            "experience": r'(?:work\s+experience|professional\s+experience|employment|career\s+history)',
            "education": r'(?:education|academic\s+background|qualifications|degrees)',
            "skills": r'(?:skills|technical\s+skills|competencies|expertise)',
            "projects": r'(?:projects|personal\s+projects|professional\s+projects)',
            "certifications": r'(?:certifications|certificates|licenses)',
            "summary": r'(?:summary|objective|profile|about)'
        }

        lines = text.split('\n')
        current_section = None
        section_content = []

        for line in lines:
            line_lower = line.lower().strip()

            # Check if line is a section header
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line_lower):
                    # Save previous section
                    if current_section and section_content:
                        sections[current_section] = '\n'.join(section_content)

                    current_section = section_name
                    section_content = []
                    break
            else:
                if current_section:
                    section_content.append(line)

        # Save last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)

        return sections

    def _extract_personal_info(self, text: str, entities: Dict[str, List[str]]) -> Dict[str, Any]:
        """Extract personal information from text and entities."""
        personal_info = {
            "name": None,
            "email": None,
            "phone": None,
            "location": None
        }

        # Extract name from entities
        if "PERSON" in entities and entities["PERSON"]:
            personal_info["name"] = entities["PERSON"][0]

        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            personal_info["email"] = email_match.group()

        # Extract phone
        phone_pattern = r'\b(?:\+?1[-.\s]?)?\(?([0-9]{3})\)?[-.\s]?([0-9]{3})[-.\s]?([0-9]{4})\b'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            personal_info["phone"] = phone_match.group()

        # Extract location
        if "GPE" in entities and entities["GPE"]:
            personal_info["location"] = entities["GPE"][0]

        return personal_info

    def _extract_work_experience(self, experience_text: str) -> List[Dict[str, Any]]:
        """Extract work experience information."""
        experiences = []

        if not experience_text:
            return experiences

        # Split by common delimiters
        experience_blocks = re.split(r'\n\s*\n|\n(?=[A-Z][^a-z]*\s*[-–—]\s*)', experience_text)

        for block in experience_blocks:
            if len(block.strip()) < 50:  # Skip very short blocks
                continue

            lines = block.strip().split('\n')
            if not lines:
                continue

            experience = {
                "title": None,
                "company": None,
                "duration": None,
                "description": []
            }

            # Try to extract title and company from first line
            first_line = lines[0].strip()
            if ' at ' in first_line or ' - ' in first_line or ' – ' in first_line:
                parts = re.split(r'\s+(?:at|-|–)\s+', first_line, 1)
                if len(parts) == 2:
                    experience["title"] = parts[0].strip()
                    experience["company"] = parts[1].strip()

            # Look for duration patterns
            duration_pattern = r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{1,2}/\d{1,2}|\d{4})\s*(?:-|to|–)\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December|\d{1,2}/\d{1,2}|\d{4}|Present|Current)'

            for line in lines:
                duration_match = re.search(duration_pattern, line)
                if duration_match:
                    experience["duration"] = duration_match.group()
                    break

            # Extract description (remaining lines)
            description_lines = []
            for line in lines[1:]:
                line = line.strip()
                if line and len(line) > 10:
                    # Clean up bullet points
                    line = re.sub(r'^[•\-\*\+\s]*', '', line)
                    description_lines.append(line)

            experience["description"] = description_lines

            if experience["title"] or experience["company"]:
                experiences.append(experience)

        return experiences

    def _extract_education(self, education_text: str) -> List[Dict[str, Any]]:
        """Extract education information."""
        education_list = []

        if not education_text:
            return education_list

        # Split into education blocks
        education_blocks = re.split(r'\n\s*\n', education_text)

        for block in education_blocks:
            lines = block.strip().split('\n')
            if not lines:
                continue

            education = {
                "degree": None,
                "institution": None,
                "year": None,
                "details": []
            }

            # Try to identify degree and institution
            for line in lines:
                line = line.strip()

                # Look for degree patterns
                degree_patterns = [
                    r'\b(?:Bachelor|Master|PhD|Doctorate|Associate|MBA|MSc|BSc|BEng|MEng|MA|BA)\b.*?(?:of|in)\s+([A-Za-z\s]+)',
                    r'\b([A-Za-z\s]+)\s+(?:Degree|Certificate|Diploma)\b'
                ]

                for pattern in degree_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        education["degree"] = match.group()
                        break

                # Look for institution
                if not education["institution"]:
                    # Common university keywords
                    if any(word in line.lower() for word in ['university', 'college', 'institute', 'school']):
                        education["institution"] = line

                # Look for year
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                if year_match:
                    education["year"] = year_match.group()

            # Extract additional details
            details = []
            for line in lines:
                line = line.strip()
                if line and len(line) > 5 and line not in [education.get("degree"), education.get("institution")]:
                    details.append(line)

            education["details"] = details

            if education["degree"] or education["institution"]:
                education_list.append(education)

        return education_list

    def _extract_highlights(self, text: str) -> List[str]:
        """Extract key highlights and achievements."""
        highlights = []

        # Look for achievement indicators
        achievement_patterns = [
            r'•\s*(.+)',  # Bullet points
            r'-\s*(.+)',  # Dash points
            r'✓\s*(.+)',  # Check marks
            r'★\s*(.+)',  # Stars
        ]

        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            for pattern in achievement_patterns:
                match = re.match(pattern, line)
                if match:
                    highlight = match.group(1).strip()
                    if len(highlight) > 10:  # Only meaningful highlights
                        highlights.append(highlight)
                    break

        # Limit to top 5 highlights
        return highlights[:5]

    async def _analyze_with_gemini(self, text: str) -> Dict[str, Any]:
        """Fallback to Google Gemini for analysis (original implementation)."""
        if not self.gemini_available:
            raise HTTPException(
                status_code=500,
                detail="No NLP models available for analysis"
            )

        try:
            prompt = f"""
            Analyze the following resume text and extract information in JSON format:

            {text}

            Please extract and categorize the following information in this exact JSON structure:
            {{
                "personalInfo": {{
                    "name": "candidate's full name",
                    "email": "email if found",
                    "phone": "phone if found",
                    "location": "location if found"
                }},
                "workExperience": [
                    {{
                        "title": "job title",
                        "company": "company name",
                        "duration": "employment period",
                        "description": ["bullet point 1", "bullet point 2"]
                    }}
                ],
                "education": [
                    {{
                        "degree": "degree name",
                        "institution": "school name",
                        "year": "graduation year",
                        "details": ["relevant detail 1", "relevant detail 2"]
                    }}
                ],
                "skills": ["skill1", "skill2", "skill3"],
                "highlights": ["achievement1", "achievement2"]
            }}

            Instructions:
            1. Include all dates in consistent format
            2. Split descriptions into clear bullet points
            3. Extract all relevant skills mentioned
            4. Keep the exact JSON structure as shown
            5. Return only the JSON object, no additional text
            """
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = await model.generate_content_async(prompt)
            response_text = response.text

            if not response_text:
                raise ValueError("No response from Gemini")

            json_str = re.search(r'({.*})', response_text, re.DOTALL)
            if not json_str:
                raise ValueError("No JSON found in response")

            result = json.loads(json_str.group(1))
            return result

        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse resume analysis result: {str(e)}"
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to analyze resume: {str(e)}"
            )
