import PyPDF2
import docx
import pdfplumber
from fastapi import UploadFile, HTTPException
import re
from typing import Dict, Any
import json
from io import BytesIO
import google.generativeai as genai


class ResumeParser:
    async def parse(self, file: UploadFile) -> Dict[str, Any]:
        """Parse resume file and extract information"""
        if not file or not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
            
        content = await self._extract_text(file)
        return await self._analyze_with_gemini(content)

    async def _extract_text(self, file: UploadFile) -> str:
        """Extract text from PDF or DOCX file"""
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
            
        try:
            content = await file.read()
            text = ""
            file_bytes = BytesIO(content)
            filename = file.filename.lower()

            if filename.endswith('.pdf'):
                # Try with PyPDF2 first
                try:
                    pdf_reader = PyPDF2.PdfReader(file_bytes)
                    text = " ".join(page.extract_text() for page in pdf_reader.pages)
                except Exception as e:
                    # If PyPDF2 fails, try with pdfplumber
                    try:
                        file_bytes.seek(0)  # Reset file pointer
                        with pdfplumber.open(file_bytes) as pdf:
                            text = " ".join(page.extract_text() for page in pdf.pages)
                    except Exception as pdf_e:
                        raise HTTPException(
                            status_code=400,
                            detail=f"Failed to read PDF file: {str(e)}. pdfplumber error: {str(pdf_e)}"                        )
            
            elif filename.endswith(('.doc', '.docx')):
                try:
                    file_bytes.seek(0)  # Reset file pointer
                    doc = docx.Document(file_bytes)
                    paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
                    if not paragraphs:
                        # If no paragraphs found, try to extract from tables as well
                        for table in doc.tables:
                            for row in table.rows:
                                paragraphs.extend(cell.text.strip() for cell in row.cells if cell.text.strip())
                    
                    text = " ".join(paragraphs)
                    
                    if not text.strip():
                        raise HTTPException(
                            status_code=400,
                            detail="No readable text found in the Word document. Please check if the document contains text content."
                        )
                except Exception as e:
                    raise HTTPException(
                        status_code=400,
                        detail=f"Failed to read Word document: {str(e)}. Please ensure the document is not corrupted and is a valid .doc or .docx file."
                    )
            
            else:
                raise HTTPException(
                    status_code=400,
                    detail="Unsupported file format. Please upload a PDF or DOCX file."
                )

            if not text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="No text could be extracted from the file."
                )
                
            return text.strip()
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process file: {str(e)}"
            )

    async def _analyze_with_gemini(self, text: str) -> Dict[str, Any]:
        """Use Google Gemini to analyze resume text"""
        try:
            prompt = f"""
            Analyze the following resume text and extract information in JSON format:
            
            {text}
            
            Please extract and categorize the following information in this exact JSON structure:
            {{
                "personalInfo": {{
                    "name": "candidate's full name",
                    "email": "email if found",
                    "phone": "phone if found",
                    "location": "location if found"
                }},
                "workExperience": [
                    {{
                        "title": "job title",
                        "company": "company name",
                        "duration": "employment period",
                        "description": ["bullet point 1", "bullet point 2"]
                    }}
                ],
                "education": [
                    {{
                        "degree": "degree name",
                        "institution": "school name",
                        "year": "graduation year",
                        "details": ["relevant detail 1", "relevant detail 2"]
                    }}
                ],
                "skills": ["skill1", "skill2", "skill3"],
                "highlights": ["achievement1", "achievement2"]
            }}
            
            Instructions:
            1. Include all dates in consistent format
            2. Split descriptions into clear bullet points
            3. Extract all relevant skills mentioned
            4. Keep the exact JSON structure as shown
            5. Return only the JSON object, no additional text
            """
              # Use the GenerativeModel directly
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = await model.generate_content_async(prompt)
            response_text = response.text

            # Extract JSON from the response
            if not response_text:
                raise ValueError("No response from Gemini")
                
            json_str = re.search(r'({.*})', response_text, re.DOTALL)
            if not json_str:
                raise ValueError("No JSON found in response")
                
            result = json.loads(json_str.group(1))
            return result
            
        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse resume analysis result: {str(e)}"
            )
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to analyze resume: {str(e)}"
            )
